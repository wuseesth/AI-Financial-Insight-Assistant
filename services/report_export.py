"""
报告导出服务模块
================
支持将分析结果导出为 PDF 和 Word 格式。
"""

import io
import os
from datetime import datetime
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable,
)
from reportlab.lib import colors


class ReportExporter:
    """报告导出服务，支持 PDF 和 Word 格式"""

    # 主题色
    PRIMARY_COLOR = "#00D4AA"
    SECONDARY_COLOR = "#1A1D27"
    TEXT_COLOR = "#1A1D27"
    ACCENT_COLOR = "#FFC107"

    @staticmethod
    def _get_type_title(analysis_type: str) -> str:
        """获取分析类型的中文标题"""
        titles = {
            "news_analysis": "财经新闻分析报告",
            "announcement_analysis": "上市公司公告分析报告",
            "hotspot_analysis": "市场热点分析报告",
            "stock_deep_decode": "股票深度解码报告",
            "stock_comparison": "多股票对比分析报告",
            "kline_chart": "K线图行情报告",
        }
        return titles.get(analysis_type, "分析报告")

    @staticmethod
    def _get_type_icon(analysis_type: str) -> str:
        """获取分析类型的图标"""
        icons = {
            "news_analysis": "📰",
            "announcement_analysis": "📋",
            "hotspot_analysis": "🔥",
            "stock_deep_decode": "🔍",
            "stock_comparison": "📊",
            "kline_chart": "📈",
        }
        return icons.get(analysis_type, "📄")

    # ============================================================
    # Word 导出
    # ============================================================

    @classmethod
    def export_to_word(cls, result: Dict[str, Any], analysis_type: str) -> bytes:
        """
        将分析结果导出为 Word 文档

        Args:
            result: 分析结果字典
            analysis_type: 分析类型

        Returns:
            Word 文档的 bytes 数据
        """
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x1A, 0x1D, 0x27)

        # 标题
        title = doc.add_heading("", level=0)
        run = title.add_run(
            f"{cls._get_type_icon(analysis_type)} {cls._get_type_title(analysis_type)}"
        )
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x00, 0xD4, 0xAA)

        # 生成时间
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x92, 0xB0)

        doc.add_paragraph("")  # 空行

        # 根据类型渲染内容
        if analysis_type == "news_analysis":
            cls._render_news_to_word(doc, result)
        elif analysis_type == "announcement_analysis":
            cls._render_announcement_to_word(doc, result)
        elif analysis_type == "hotspot_analysis":
            cls._render_hotspot_to_word(doc, result)
        elif analysis_type == "stock_deep_decode":
            cls._render_stock_decode_to_word(doc, result)
        elif analysis_type == "stock_comparison":
            cls._render_comparison_to_word(doc, result)
        elif analysis_type == "kline_chart":
            cls._render_kline_to_word(doc, result)
        else:
            doc.add_paragraph("暂不支持该类型的报告导出")

        # 添加免责声明
        doc.add_paragraph("")
        doc.add_paragraph("—" * 50)
        disclaimer = doc.add_paragraph()
        run = disclaimer.add_run(
            "免责声明：本报告由 AI 自动生成，仅供参考，不构成投资建议。投资有风险，决策需谨慎。"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x92, 0xB0)
        run.italic = True

        # 保存到 BytesIO
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    @classmethod
    def _add_section_title(cls, doc: Document, text: str):
        """添加章节标题"""
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x00, 0xD4, 0xAA)

    @classmethod
    def _add_subsection_title(cls, doc: Document, text: str):
        """添加子章节标题"""
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1D, 0x27)

    @classmethod
    def _add_key_value(cls, doc: Document, key: str, value: str):
        """添加键值对"""
        p = doc.add_paragraph()
        run = p.add_run(f"{key}：")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(str(value))

    @classmethod
    def _render_news_to_word(cls, doc: Document, result: Dict[str, Any]):
        """渲染新闻分析到 Word"""
        cls._add_section_title(doc, "📝 新闻总结")
        doc.add_paragraph(result.get("summary", "无"))

        sentiment = result.get("sentiment", {})
        if sentiment:
            cls._add_section_title(doc, "📊 情绪分析")
            cls._add_key_value(doc, "情绪判断", sentiment.get("judgment", "中性"))
            cls._add_key_value(doc, "置信度", f"{sentiment.get('confidence', 0)}%")
            if sentiment.get("reason"):
                cls._add_key_value(doc, "判断依据", sentiment["reason"])

        if result.get("recommendation"):
            cls._add_section_title(doc, "💡 投资建议")
            doc.add_paragraph(result["recommendation"])

        industries = result.get("industries", [])
        if industries:
            cls._add_section_title(doc, "🏭 涉及行业")
            doc.add_paragraph("、".join(industries))

        companies = result.get("companies", [])
        if companies:
            cls._add_section_title(doc, "🏢 涉及公司")
            doc.add_paragraph("、".join(companies))

        risks = result.get("risks", [])
        if risks:
            cls._add_section_title(doc, "⚠️ 风险提示")
            for risk in risks:
                doc.add_paragraph(f"• {risk}")

    @classmethod
    def _render_announcement_to_word(cls, doc: Document, result: Dict[str, Any]):
        """渲染公告分析到 Word"""
        cls._add_key_value(doc, "公告标题", result.get("title", "无"))
        cls._add_key_value(doc, "公司名称", result.get("company", "无"))
        cls._add_key_value(doc, "公告日期", result.get("date", "无"))

        cls._add_section_title(doc, "📌 核心事件")
        doc.add_paragraph(result.get("core_event", "无"))

        cls._add_section_title(doc, "💰 财务数据分析")
        doc.add_paragraph(result.get("financial_analysis", "无"))

        cls._add_section_title(doc, "📈 市场影响")
        doc.add_paragraph(result.get("market_impact", "无"))

        cls._add_section_title(doc, "💡 投资建议")
        doc.add_paragraph(result.get("investment_advice", "无"))

    @classmethod
    def _render_hotspot_to_word(cls, doc: Document, result: Dict[str, Any]):
        """渲染热点分析到 Word"""
        cls._add_section_title(doc, "🔥 热点话题")
        doc.add_paragraph(result.get("hotspot_topic", "无"))

        cls._add_section_title(doc, "📊 热门行业")
        industries = result.get("hot_industries", [])
        if industries:
            for ind in industries:
                name = ind.get("name", "") if isinstance(ind, dict) else ind
                reason = ind.get("reason", "") if isinstance(ind, dict) else ""
                doc.add_paragraph(f"• {name}")
                if reason:
                    doc.add_paragraph(f"  原因：{reason}")

        cls._add_section_title(doc, "🏢 热门公司")
        companies = result.get("hot_companies", [])
        if companies:
            for comp in companies:
                name = comp.get("name", "") if isinstance(comp, dict) else comp
                reason = comp.get("reason", "") if isinstance(comp, dict) else ""
                doc.add_paragraph(f"• {name}")
                if reason:
                    doc.add_paragraph(f"  原因：{reason}")

        if result.get("market_analysis"):
            cls._add_section_title(doc, "📈 市场分析")
            doc.add_paragraph(result["market_analysis"])

    @classmethod
    def _render_stock_decode_to_word(cls, doc: Document, result: Dict[str, Any]):
        """渲染股票深度解码到 Word"""
        # 基本信息
        basic = result.get("basic_info", {})
        if basic:
            cls._add_section_title(doc, "📋 基本信息")
            for key, value in basic.items():
                cls._add_key_value(doc, key, str(value))

        # 市场归属与交易规则
        market = result.get("market_rules", {})
        if market:
            cls._add_section_title(doc, "🌐 市场归属与交易规则")
            for key, value in market.items():
                cls._add_key_value(doc, key, str(value))

        # 异动分析
        anomaly = result.get("anomaly_analysis", {})
        if anomaly:
            cls._add_section_title(doc, "⚡ 异动分析")
            for key, value in anomaly.items():
                cls._add_key_value(doc, key, str(value))

        # 资金行为
        capital = result.get("capital_behavior", {})
        if capital:
            cls._add_section_title(doc, "💰 资金行为分析")
            for key, value in capital.items():
                cls._add_key_value(doc, key, str(value))

        # 策略建议
        strategy = result.get("strategy_advice", {})
        if strategy:
            cls._add_section_title(doc, "💡 策略建议")
            for key, value in strategy.items():
                cls._add_key_value(doc, key, str(value))

    @classmethod
    def _render_comparison_to_word(cls, doc: Document, result: Dict[str, Any]):
        """渲染多股票对比到 Word"""
        # 对比表格
        table_data = result.get("comparison_table", [])
        if table_data and isinstance(table_data, list) and len(table_data) > 0:
            cls._add_section_title(doc, "📊 对比表格")
            headers = list(table_data[0].keys()) if isinstance(table_data[0], dict) else []
            if headers:
                rows = [headers]
                for row_data in table_data:
                    if isinstance(row_data, dict):
                        rows.append([str(row_data.get(h, "")) for h in headers])
                table = doc.add_table(rows=len(rows), cols=len(headers))
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, row_data in enumerate(rows):
                    for j, cell_text in enumerate(row_data):
                        table.cell(i, j).text = cell_text

        # 关键差异
        diffs = result.get("key_differences", "")
        if diffs:
            cls._add_section_title(doc, "🔍 关键差异")
            doc.add_paragraph(diffs)

        # 跨市场套利
        arbitrage = result.get("cross_market_arbitrage", "")
        if arbitrage:
            cls._add_section_title(doc, "🔄 跨市场套利机会")
            doc.add_paragraph(arbitrage)

        # 投资策略
        strategy = result.get("investment_strategy", "")
        if strategy:
            cls._add_section_title(doc, "💡 投资策略")
            doc.add_paragraph(strategy)

        # 风险提示
        risks = result.get("risk_warnings", "")
        if risks:
            cls._add_section_title(doc, "⚠️ 风险提示")
            doc.add_paragraph(risks)

    @classmethod
    def _render_kline_to_word(cls, doc: Document, result: Dict[str, Any]):
        """渲染K线图到 Word"""
        symbol = result.get("symbol", "")
        if symbol:
            cls._add_key_value(doc, "股票代码", symbol)

        info = result.get("info", {})
        if info:
            cls._add_key_value(doc, "股票名称", info.get("name", ""))
            cls._add_key_value(doc, "市场", info.get("market", ""))
            cls._add_key_value(doc, "最新价", str(info.get("price", "")))
            cls._add_key_value(doc, "涨跌幅", str(info.get("change", "")))

        stats = result.get("stats", {})
        if stats:
            cls._add_section_title(doc, "📊 数据统计")
            cls._add_key_value(doc, "区间最高", f"{stats.get('high', '—')}")
            cls._add_key_value(doc, "区间最低", f"{stats.get('low', '—')}")
            cls._add_key_value(doc, "区间均价", f"{stats.get('avg', '—')}")
            cls._add_key_value(doc, "区间涨跌幅", f"{stats.get('change', '—')}%")
            cls._add_key_value(doc, "最新收盘", f"{stats.get('close', '—')}")

        doc.add_paragraph("")
        doc.add_paragraph(
            "📈 K线图无法在 Word 文档中直接显示，请在应用中查看交互式图表。"
        )

    # ============================================================
    # PDF 导出
    # ============================================================

    @classmethod
    def export_to_pdf(cls, result: Dict[str, Any], analysis_type: str) -> bytes:
        """
        将分析结果导出为 PDF 文档

        Args:
            result: 分析结果字典
            analysis_type: 分析类型

        Returns:
            PDF 文档的 bytes 数据
        """
        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
            leftMargin=2.5 * cm,
            rightMargin=2.5 * cm,
        )

        styles = getSampleStyleSheet()
        story = []

        # 自定义样式
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=22,
            textColor=HexColor(cls.PRIMARY_COLOR),
            spaceAfter=6,
            alignment=1,  # center
        )
        heading1_style = ParagraphStyle(
            "CustomH1",
            parent=styles["Heading1"],
            fontSize=16,
            textColor=HexColor(cls.PRIMARY_COLOR),
            spaceBefore=16,
            spaceAfter=8,
        )
        heading2_style = ParagraphStyle(
            "CustomH2",
            parent=styles["Heading2"],
            fontSize=13,
            textColor=HexColor(cls.TEXT_COLOR),
            spaceBefore=12,
            spaceAfter=6,
        )
        normal_style = ParagraphStyle(
            "CustomNormal",
            parent=styles["Normal"],
            fontSize=10,
            textColor=HexColor(cls.TEXT_COLOR),
            spaceAfter=4,
            leading=16,
        )
        meta_style = ParagraphStyle(
            "MetaStyle",
            parent=styles["Normal"],
            fontSize=8,
            textColor=HexColor("#8892B0"),
            spaceAfter=12,
        )
        disclaimer_style = ParagraphStyle(
            "Disclaimer",
            parent=styles["Normal"],
            fontSize=8,
            textColor=HexColor("#8892B0"),
            italic=True,
            spaceBefore=20,
        )

        # 标题
        story.append(Paragraph(
            f"{cls._get_type_icon(analysis_type)} {cls._get_type_title(analysis_type)}",
            title_style,
        ))
        story.append(Paragraph(
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            meta_style,
        ))
        story.append(HRFlowable(
            width="100%",
            thickness=1,
            color=HexColor(cls.PRIMARY_COLOR),
        ))
        story.append(Spacer(1, 12))

        # 根据类型渲染内容
        if analysis_type == "news_analysis":
            cls._render_news_to_pdf(story, result, heading1_style, heading2_style, normal_style)
        elif analysis_type == "announcement_analysis":
            cls._render_announcement_to_pdf(story, result, heading1_style, heading2_style, normal_style)
        elif analysis_type == "hotspot_analysis":
            cls._render_hotspot_to_pdf(story, result, heading1_style, heading2_style, normal_style)
        elif analysis_type == "stock_deep_decode":
            cls._render_stock_decode_to_pdf(story, result, heading1_style, heading2_style, normal_style)
        elif analysis_type == "stock_comparison":
            cls._render_comparison_to_pdf(story, result, heading1_style, heading2_style, normal_style)
        elif analysis_type == "kline_chart":
            cls._render_kline_to_pdf(story, result, heading1_style, heading2_style, normal_style)
        else:
            story.append(Paragraph("暂不支持该类型的报告导出", normal_style))

        # 免责声明
        story.append(Spacer(1, 20))
        story.append(HRFlowable(
            width="100%",
            thickness=0.5,
            color=HexColor("#CCCCCC"),
        ))
        story.append(Paragraph(
            "免责声明：本报告由 AI 自动生成，仅供参考，不构成投资建议。投资有风险，决策需谨慎。",
            disclaimer_style,
        ))

        doc.build(story)
        buf.seek(0)
        return buf.getvalue()

    @classmethod
    def _add_pdf_section(cls, story, title: str, heading_style, content: str, normal_style):
        """添加 PDF 章节"""
        story.append(Paragraph(title, heading_style))
        story.append(Paragraph(content.replace("\n", "<br/>"), normal_style))

    @classmethod
    def _render_news_to_pdf(cls, story, result, h1, h2, normal):
        """渲染新闻分析到 PDF"""
        cls._add_pdf_section(story, "📝 新闻总结", h1, result.get("summary", "无"), normal)

        sentiment = result.get("sentiment", {})
        if sentiment:
            cls._add_pdf_section(story, "📊 情绪分析", h1, "", normal)
            story.append(Paragraph(
                f"<b>情绪判断：</b>{sentiment.get('judgment', '中性')}",
                normal,
            ))
            story.append(Paragraph(
                f"<b>置信度：</b>{sentiment.get('confidence', 0)}%",
                normal,
            ))
            if sentiment.get("reason"):
                story.append(Paragraph(
                    f"<b>判断依据：</b>{sentiment['reason']}",
                    normal,
                ))

        if result.get("recommendation"):
            cls._add_pdf_section(story, "💡 投资建议", h1, result["recommendation"], normal)

        industries = result.get("industries", [])
        if industries:
            cls._add_pdf_section(story, "🏭 涉及行业", h1, "、".join(industries), normal)

        companies = result.get("companies", [])
        if companies:
            cls._add_pdf_section(story, "🏢 涉及公司", h1, "、".join(companies), normal)

        risks = result.get("risks", [])
        if risks:
            cls._add_pdf_section(story, "⚠️ 风险提示", h1, "", normal)
            for risk in risks:
                story.append(Paragraph(f"• {risk}", normal))

    @classmethod
    def _render_announcement_to_pdf(cls, story, result, h1, h2, normal):
        """渲染公告分析到 PDF"""
        story.append(Paragraph(f"<b>公告标题：</b>{result.get('title', '无')}", normal))
        story.append(Paragraph(f"<b>公司名称：</b>{result.get('company', '无')}", normal))
        story.append(Paragraph(f"<b>公告日期：</b>{result.get('date', '无')}", normal))

        cls._add_pdf_section(story, "📌 核心事件", h1, result.get("core_event", "无"), normal)
        cls._add_pdf_section(story, "💰 财务数据分析", h1, result.get("financial_analysis", "无"), normal)
        cls._add_pdf_section(story, "📈 市场影响", h1, result.get("market_impact", "无"), normal)
        cls._add_pdf_section(story, "💡 投资建议", h1, result.get("investment_advice", "无"), normal)

    @classmethod
    def _render_hotspot_to_pdf(cls, story, result, h1, h2, normal):
        """渲染热点分析到 PDF"""
        cls._add_pdf_section(story, "🔥 热点话题", h1, result.get("hotspot_topic", "无"), normal)

        industries = result.get("hot_industries", [])
        if industries:
            cls._add_pdf_section(story, "📊 热门行业", h1, "", normal)
            for ind in industries:
                name = ind.get("name", "") if isinstance(ind, dict) else ind
                reason = ind.get("reason", "") if isinstance(ind, dict) else ""
                text = f"• {name}"
                if reason:
                    text += f"<br/>&nbsp;&nbsp;原因：{reason}"
                story.append(Paragraph(text, normal))

        companies = result.get("hot_companies", [])
        if companies:
            cls._add_pdf_section(story, "🏢 热门公司", h1, "", normal)
            for comp in companies:
                name = comp.get("name", "") if isinstance(comp, dict) else comp
                reason = comp.get("reason", "") if isinstance(comp, dict) else ""
                text = f"• {name}"
                if reason:
                    text += f"<br/>&nbsp;&nbsp;原因：{reason}"
                story.append(Paragraph(text, normal))

        if result.get("market_analysis"):
            cls._add_pdf_section(story, "📈 市场分析", h1, result["market_analysis"], normal)

    @classmethod
    def _render_stock_decode_to_pdf(cls, story, result, h1, h2, normal):
        """渲染股票深度解码到 PDF"""
        sections = [
            ("basic_info", "📋 基本信息"),
            ("market_rules", "🌐 市场归属与交易规则"),
            ("anomaly_analysis", "⚡ 异动分析"),
            ("capital_behavior", "💰 资金行为分析"),
            ("strategy_advice", "💡 策略建议"),
        ]
        for key, title in sections:
            data = result.get(key, {})
            if data:
                cls._add_pdf_section(story, title, h1, "", normal)
                if isinstance(data, dict):
                    for k, v in data.items():
                        story.append(Paragraph(f"<b>{k}：</b>{v}", normal))
                else:
                    story.append(Paragraph(str(data), normal))

    @classmethod
    def _render_comparison_to_pdf(cls, story, result, h1, h2, normal):
        """渲染多股票对比到 PDF"""
        table_data = result.get("comparison_table", [])
        if table_data and isinstance(table_data, list) and len(table_data) > 0:
            cls._add_pdf_section(story, "📊 对比表格", h1, "", normal)
            if isinstance(table_data[0], dict):
                headers = list(table_data[0].keys())
                rows = [[str(r.get(h, "")) for h in headers] for r in table_data]
                data = [headers] + rows
                col_widths = [A4[0] / len(headers) - 2 * cm] * len(headers)
                table = Table(data, colWidths=col_widths, repeatRows=1)
                table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), HexColor(cls.PRIMARY_COLOR)),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HexColor("#F5F5F5")]),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(table)
                story.append(Spacer(1, 8))

        for key, title in [
            ("key_differences", "🔍 关键差异"),
            ("cross_market_arbitrage", "🔄 跨市场套利机会"),
            ("investment_strategy", "💡 投资策略"),
            ("risk_warnings", "⚠️ 风险提示"),
        ]:
            content = result.get(key, "")
            if content:
                cls._add_pdf_section(story, title, h1, content, normal)

    @classmethod
    def _render_kline_to_pdf(cls, story, result, h1, h2, normal):
        """渲染K线图到 PDF"""
        symbol = result.get("symbol", "")
        if symbol:
            story.append(Paragraph(f"<b>股票代码：</b>{symbol}", normal))

        info = result.get("info", {})
        if info:
            story.append(Paragraph(f"<b>股票名称：</b>{info.get('name', '')}", normal))
            story.append(Paragraph(f"<b>市场：</b>{info.get('market', '')}", normal))
            story.append(Paragraph(f"<b>最新价：</b>{info.get('price', '')}", normal))
            story.append(Paragraph(f"<b>涨跌幅：</b>{info.get('change', '')}", normal))

        stats = result.get("stats", {})
        if stats:
            cls._add_pdf_section(story, "📊 数据统计", h1, "", normal)
            story.append(Paragraph(f"<b>区间最高：</b>{stats.get('high', '—')}", normal))
            story.append(Paragraph(f"<b>区间最低：</b>{stats.get('low', '—')}", normal))
            story.append(Paragraph(f"<b>区间均价：</b>{stats.get('avg', '—')}", normal))
            story.append(Paragraph(f"<b>区间涨跌幅：</b>{stats.get('change', '—')}%", normal))
            story.append(Paragraph(f"<b>最新收盘：</b>{stats.get('close', '—')}", normal))

        story.append(Paragraph(
            "📈 K线图无法在 PDF 中直接显示，请在应用中查看交互式图表。",
            normal,
        ))
